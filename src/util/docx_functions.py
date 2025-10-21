from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from io import BytesIO
from docx.shared import Inches
import re


class DocumentHandler:
    @staticmethod
    def create_docx_from_txt(input_file, output_file):
        doc = Document()
        heading_style = doc.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.size = Pt(16)
        heading_style.font.bold = True

        with open(input_file, 'r', encoding='utf-8') as file:
            lines = file.readlines()

        for line in lines:
            line = line.strip()
            if line.startswith('###'):
                doc.add_paragraph(line.strip('### '), style='CustomHeading')
            elif '**' in line:
                p = doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:
                        run.bold = True
            else:
                doc.add_paragraph(line)

        doc.save(output_file)
        print(f"DOCX file created at {output_file}")

    # Function to generate a Word document from text
    @staticmethod
    def generate_word_document(text, image_paths=None):
        doc = Document()
        
        paragraphs = text.split('\n')
        for paragraph in paragraphs:
            stripped = paragraph.strip()
            
            # Handle headings
            if stripped.startswith('#'):
                heading_level = len(stripped.split()[0])  # Count the number of '#' symbols
                heading_text = stripped.lstrip('#').strip()
                doc.add_heading(heading_text, level=heading_level)
            else:
                p = doc.add_paragraph()
                parts = re.split(r'(<SCREENSHOT>.*?</SCREENSHOT>|\*\*.*?\*\*)', paragraph)
                for part in parts:
                    if part.startswith('<SCREENSHOT>') and part.endswith('</SCREENSHOT>'):
                        if image_paths and part in image_paths:
                            run = p.add_run()
                            run.add_picture(image_paths[part], width=Inches(6))
                    elif part.startswith('**') and part.endswith('**'):
                        p.add_run(part.strip('**')).bold = True
                    else:
                        p.add_run(part)
        
        # Save the document to a binary object
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        
        return bio

    @staticmethod
    def extract_text_from_docx(file_path):
        try:
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n'.join(full_text)
        except Exception as e:
            print(f"An error occurred: {e}")
            return None